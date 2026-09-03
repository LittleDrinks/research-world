"""按官方模板填充 P1–P20 技术方案报告。

用法: /path/to/python fill_report.py（仓库根 research-world/ 下运行）
输入: docs/赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx
输出: evidence/contest-2026/report/技术方案报告-ResearchWorld-filled.docx

原则：P13–P19 数字与 evidence/contest-2026/submission-evidence.md 完全一致；
团队信息/报名截图/部署 URL/视频链接留【待用户补充】占位，不编造。
"""
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[4]  # research-world/
TEMPLATE = ROOT / "docs" / "赛道一-方向1A-科学假设生成与研究计划设计-提交要求及模板.docx"
FIGDIR = Path(__file__).resolve().parent / "figures"
OUT = Path(__file__).resolve().parent / "技术方案报告-ResearchWorld-filled.docx"

PENDING = "【待用户补充"


def fmt(run, size=10.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")


def fill_cell(cell, text, size=8.5):
    cell.text = ""
    for i, part in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        fmt(p.add_run(part), size)


def fill_table(doc, idx, start_row, col_start, rows, size=8.5):
    t = doc.tables[idx]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            fill_cell(t.rows[start_row + ri].cells[col_start + ci], val, size)


def replace_para(doc, prefix, text, size=10.5, bold=False, keep_bullet=None):
    """按前缀找段落并整体替换文本。keep_bullet=True 时保留原前导项目符号词。"""
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            full = (keep_bullet + text) if keep_bullet else text
            fmt(p.add_run(full), size, bold)
            return True
    raise KeyError(f"paragraph not found: {prefix[:30]}")


def clear_underscore_lines(doc):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and set(t) == {"_"}:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)


def insert_figure(doc, prefix, png, caption, width=5.7):
    """把占位段替换为居中图片，并在其后插入图题段。"""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(FIGDIR / png), width=Inches(width))
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn("w:keepNext")) is None:
                pPr.append(pPr.makeelement(qn("w:keepNext"), {}))
            newp = OxmlElement("w:p")
            p._element.addnext(newp)
            capp = Paragraph(newp, p._parent)
            capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt(capp.add_run(caption), 9, bold=True)
            return True
    raise KeyError(f"figure placeholder not found: {prefix[:30]}")


def insert_figure_after_table(doc, table_idx, png, caption, width=5.7):
    """在指定表格之后插入居中图片＋图题。"""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    t = doc.tables[table_idx]
    capp_el = OxmlElement("w:p")
    t._element.addnext(capp_el)
    img_el = OxmlElement("w:p")
    t._element.addnext(img_el)  # 先插图题再插图，最终顺序：表→图→图题
    imgp = Paragraph(img_el, t._parent)
    imgp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    imgp.add_run().add_picture(str(FIGDIR / png), width=Inches(width))
    _ippr = imgp._element.get_or_add_pPr()
    if _ippr.find(qn("w:keepNext")) is None:
        _ippr.append(_ippr.makeelement(qn("w:keepNext"), {}))
    capp = Paragraph(capp_el, t._parent)
    capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(capp.add_run(caption), 9, bold=True)


# ---------------------------------------------------------------------------
# 正文文案
# ---------------------------------------------------------------------------

INTRO = (
    "本作品面向科学假设生成与研究计划设计：每题由独立 Qwen 作者 Session 校正题干前提、识别缺口、核验证据、生成三个可区分方向、"
    "比较筛选并形成含 planned/executed 边界的研究计划；独立评审以六维 rubric（满分12）与来源门逐条核验引用，findings 驱动多轮修订。"
    "已完成官方125题全量：候选结论125/125，completed 8、partial 117、failed 0，2592次调用；"
    "五深度案例最终均12/12、来源门全过；旗舰 q049 由 V1 的 9/12 修订至 V8 的 12/12、来源 6/6。账本与哈希可回读，候选结论不等同已验证发现。"
)

QWEN_NOTE = (
    "本作品以 Qwen 系列为基座，经阿里云百炼 OpenAI 兼容端点调用：作者 Session 用 qwen3-max 完成问题理解、证据检索、三方向生成、比较筛选与研究计划；"
    "评审与回执用 gpt-5.6-sol、qwen3.7-max 等交叉执行，避免自评偏置。上下文按层拼接系统协议、canonical 原题、来源记录、历史 findings 与"
    " planned/executed 边界；评审侧仅见固定 rubric 与被审版本，引用独立检索核验。verdict=revise 时 findings 逐条回流触发修订，形成多轮闭环；"
    "调用、token 与输出哈希均入账本可审计。"
)

CLAIMS = [
    ("本作品针对的具体问题：",
     "大模型单轮回答科学问题时，易把候选解释写成已证实结论：来源不可回读、无多候选比较、研究计划空泛、planned 与 executed 混写、失败不留痕。"),
    ("本作品实际完成的核心方法：",
     "「作者—独立评审」双角色多轮工作流：证据门槛前置、三方向候选假设、六维 rubric＋来源门评审、findings 驱动修订、限定计算与 planned/executed 分离、回执复算。"),
    ("最有代表性的结果1（对象、口径、结果）：",
     "q049（行星轨道不坠日）：V1 9/12、来源 2/5 → V8 12/12、来源 6/6；唯一执行项 Peters 计算 P=196.291 W、t=1.069e23 yr，评审独立复算一致。"),
    ("最有代表性的结果2（对象、口径、结果）：",
     "官方 125 题全量：候选结论 125/125，completed 8、partial 117、waiting_human 0、failed 0；五个深度案例最终均 12/12、来源门 6/6–9/9 全过。"),
    ("目前仍存在的主要局限：",
     "117 题为 partial（至少一项未过轻量门槛）；未做同题重复运行；付费全文、受限数据、伦理审批与领域阈值题不可覆盖；候选结论≠已验证发现。"),
]

P2_PROBLEM = (
    "现有方法直接以单轮问答处理科学问题，存在四项不足：（1）候选解释与已证实结论不分，无来源核验；"
    "（2）单假设输出，无可区分方向间的比较与筛选；（3）研究计划空泛，planned 与 executed 混写；"
    "（4）失败、暂停与修订过程不留痕，结论无法回查。本作品针对并解决全部四项，核心目标是让每个假设具备证据闭包：可检验、可回读、可审计。"
)

P2_WORK = (
    "已实现：（1）官方 125 题全量轻量协议运行，共 135 个 Session（125 作者＋5 stopped＋5 审计）、2592 次调用；"
    "（2）q049/q089/q021/q112/q098 五题深度运行：独立评审、多轮修订、回执复算、同条件直接回答对照，最终均 12/12 且来源门全过；"
    "（3）q049 限定计算（Peters 公式）真实执行并留哈希；（4）Research World 平台（web 前端、control 控制面、Agent Runtime、"
    "runner-controller 工具沙箱，docker compose 一键启动）与 T2 Playwright 5 项验收。规划中：Research Graph 与主 Agent 动态 Workflow（ADR-0037）。"
)

P2_BULLETS = [
    ("本作品最终形成的主要输出：",
     "每题候选结论＋研究计划＋来源记录＋运行账本；125 份逐题文档、5 份深度案例链（全部版本、评审、回执、对照与哈希保留）。"),
    ("本作品已使用官方125道科学问题完成的实际测试：",
     "125/125 全量覆盖：135 Session、2592 次调用、非缓存输入 27,871,295 token、输出 581,255 token；completed 8、partial 117、failed 0。"),
    ("相较直接使用大模型回答，本作品最主要的不同：",
     "双角色分离＋引用逐条回读＋判据可执行＋失败留痕；q049 同条件直接回答对照仅 4/12 与 6/12，均无来源、无方向、无计划。"),
]

P3_TABLE = [
    ["三方向均须标注依据来源与推断边界；证据门槛要求同行评议或机构一手来源，精确主张须有可回读标识符。",
     "防止把模型推断当文献事实；来源可回读才能支撑或否定假设。"],
    ["从「争议/未知」清单生成机制层面可区分的三个 Direction，缺口与方向一一映射。",
     "缺口可处理，假设才可检验；避免凭空生成。"],
    ["每方向含可区分预测与判定判据，判据须经独立评审验证物理正确性（如 inspiral 时间与太阳寿命比较）。",
     "判据错误会把真实约 200 W 的耗散误判为显著，导致计划失真。"],
    ["三方向并列＋横向比较＋评审记录被拒原因；反对证据与替代解释显式入文。",
     "防止单主线确认偏误；被拒方向保留可复评。"],
    ["planned/executed 显式分离；结论限定表述（如「耗散时间远超太阳寿命，轨道不会因该机制螺旋坠日」）。",
     "候选假设只是可提交、可回读的候选，不是新科学发现。"],
]

P3_LOGIC = (
    "以 q049 为例：（1）题干信息用于前提校正与知识缺口识别——「轨道衰变终将坠日」被校正为保守动力学、微弱耗散、混沌失稳与太阳演化四层机制问题；"
    "（2）同行评议一手来源用于约束假设——来源须有 DOI/arXiv 标识符且经评审逐条回读；（3）支持条件为评审 rubric 得分与来源门通过，"
    "削弱条件为引用抽查 fail 或判据无效，否定条件为复算推翻关键数值（引力波功率错 22 个数量级即被推翻重写）。"
)

P4_TABLE = [
    ["同行评议论文与机构一手资料",
     "anysearch 检索与提取；记录作者/年份/标题/DOI|arXiv 号；评审逐条回读",
     "承载机制主张、定量数值与判据（Peters 功率、Lyapunov 时间等）",
     "付费墙与不可提取的 PDF 正文记为 unverifiable"],
    ["官方 125 道 canonical 科学问题（docs/questions.json）",
     "赛事提供，作唯一输入口径",
     "统一题目身份，输出按题入库（q001–q125.md）",
     "英文原题，理解存在转写噪声"],
    ["运行账本（run.md / index.md / audit-*.md）",
     "Pi Session 自动留痕＋SHA-256 内容哈希",
     "审计 Session/调用/token/终态，支撑回执复算",
     "raw JSONL 未全部随仓库公开（脱敏副本提交时附）"],
    ["独立评审与回执记录（review-*.md / receipt-*.md）",
     "独立 Session 固定 rubric 评审＋引用核验＋账本复算",
     "形成 revise/deliverable 判定与 findings 清单",
     "评审范围限于声明读取的文件；非领域专家，需人工 Gate 兜底"],
]

P4_BULLETS = [
    ("本作品如何判断来源是否可用：",
     "优先同行评议与机构一手来源；标识符（DOI/arXiv）必须可回读；关键数值经独立核验后方可引用。"),
    ("遇到证据不足或相互冲突时，系统实际如何处理：",
     "记入「争议/未知」并保留两说；比较对象不等价即拒绝对照（attempt 2 只近似计算量、attempt 6 只近似长度）；错源删除或修正（q049 来源 4/5）。"),
    ("本作品如何保留引用和科学证据之间的对应关系：",
     "主张内联标注来源编号；来源记录作者/年份/标题/DOI；评审核验结论（pass/fail＋核验 URL）随评审文档入库。"),
]

P5_TABLE = [
    ["六维 rubric 评分",
     "独立评审 Session 固定 rubric：问题理解/文献证据/Direction/科学推理/研究计划/表达与追溯，各 0–2 分",
     "五深度案例各版本＋直接回答对照",
     "＜10 分或关键引用 fail 即 revise；达标且来源门过方可交付"],
    ["来源门（引用有效率）",
     "逐条回读 DOI/arXiv，核标识符、题名、作者与转述方向",
     "全部显式来源（V1 分母 5 条、V8 分母 6 条）",
     "门未过不得交付；错源删除或修正后重审"],
    ["回执复算",
     "独立 Session 重算模型、调用、token、写入与哈希",
     "深度案例最终版",
     "与账本不一致即拒绝（review-v15 曾被独立回执否决）"],
    ["全量轻量审计",
     "5 个独立审计 Session 各复核连续 25 题",
     "q001–q125 全量",
     "终态与缺口全保留，不以深度案例替代全量"],
]

P5_BULLETS = [
    ("团队重点评价的是哪些方面，为什么：",
     "证据可回读与 planned/executed 边界——对照实验显示直接回答的错误集中于错引与数量级，这两项直接决定结论能否回查。"),
    ("评价由程序、模型、团队成员或领域专家中的哪些主体完成：",
     "独立模型评审（与作者不同模型交叉）＋程序哈希核验＋人工 Gate（伦理/招募/阈值类题）；无领域专家常驻。"),
    ("第一轮与第二轮是否采用同一口径；如有变化，原因是什么：",
     "同一 rubric 与来源门；变化仅在评审读取范围扩大（终审复核 v7＋v8）并加入回执复算。"),
    ("什么情况下继续生成、停止或交由研究者判断：",
     "verdict=revise 继续；deliverable＋回执一致＋剩余项全部 planned 即停止；伦理审批、付费数据或领域阈值题交 waiting_human。"),
]

P6_TABLE = [
    ["canonical 原题＋题干",
     "前提校正；对象/范围/关键变量提取；已有认识、争议与未知三分",
     "知识缺口清单",
     "缺口直接决定三方向的生成空间"],
    ["知识缺口＋检索预算",
     "anysearch 检索/提取；来源记录（作者/年份/标题/DOI）；支持与反对证据分列",
     "带标识符的来源列表",
     "约束方向取舍与研究计划判据"],
    ["缺口＋证据",
     "生成三个机制层面可区分 Direction＋差异维度声明＋横向比较",
     "D1–D3 候选集",
     "整体送独立评审"],
    ["D1–D3＋来源",
     "六维 rubric 打分＋引用逐条回读＋与直接回答基线对照",
     "verdict＋findings 清单",
     "findings 逐条回流触发修订"],
    ["选定方向＋findings",
     "研究步骤/判据/planned 标注/限定计算＋回执复算",
     "最终版＋研究计划＋账本",
     "计入运行账本，支撑全量审计"],
]

P6_LOOP = (
    "本作品不是一次性问答：评审 findings 是显式中间结果，逐条进入下一轮作者上下文——q049 review-v1 的 5 条 findings（错配 DOI、错 arXiv 号、"
    "反向转述、22 个数量级错误、无效判据）在后续版本中逐条处置；版本链保留每轮被拒原因，回执复算防止「自己证明自己」；"
    "终态机（completed/partial/waiting_human/failed）把失败与人工暂停全部留痕。规划中的 Research Graph（ADR-0037）将把 Direction/Experiment/Claim "
    "的证据闭包写入图谱，报告只读取有证据闭包的 Claim。"
)

P7_TABLE = [
    ["作者/基线：qwen3-max；评审/回执：gpt-5.6-sol、qwen3.7-max 等。经阿里云百炼 OpenAI 兼容端点，由 Agent Runtime 以独立 Session 调用，凭证隔离（.env 不入库）。"],
    ["作者：前提校正、缺口识别、证据检索、三方向生成、比较筛选、研究计划与候选结论；评审：rubric 打分＋引用核验；回执：账本独立复算。"],
    ["系统协议（输出契约/证据门槛/文件边界/检索预算）、canonical 原题、来源记录、历史版本与评审 findings、planned/executed 约束（见图3）。"],
    ["固定章节 Markdown 报告（研究对象/已有认识/来源记录/D1–D3/研究计划/planned 标注）；run.md frontmatter 记录终态、模型与哈希。"],
    ["web_search/extract 检索、仓库文件写入、Peters 公式脚本执行、SHA-256 计算；评审侧只读被审文档并独立检索核验。"],
]

P7_BULLETS = [
    ("团队为减少无依据生成实际采取的措施：",
     "证据门槛前置声明；「未直接检索到不得引用」；来源须有可回读标识符；评审用独立检索核验而非采信作者记录。"),
    ("上下文更新或多轮调用的触发方式：",
     "verdict=revise 即开新一轮并把 findings 逐条注入；deliverable 且回执一致即停；失败 Session 以全新 Session 重试并留痕。"),
    ("该设计对结果带来的实际变化：",
     "q049 V1 9/12（来源 2/5）→V8 12/12（6/6）；错引、反向转述、22 个数量级错误与无效判据逐轮清零。"),
]

P8_TABLE = [
    ["题干解析＋前提校正",
     "研究对象定为太阳系八大行星轨道长期稳定性；范围至太阳主序星结束",
     "排除系外对比类内容，聚焦机制比较"],
    ["从对象与机制清单提取",
     "轨道要素 a/e/i、潮汐品质因子 Q、引力波功率、太阳质量损失率",
     "构成计划判据与限定计算的输入"],
    ["检索＋整理",
     "已有：Gyr 尺度数值稳定；争议：水星失稳概率 1–2% 或更高；未知：外行星 ＞10 Gyr 稳定性",
     "缺口取自争议与未知的交集"],
    ["汇总前三步",
     "机制时间尺度对比、混沌与耗散相对重要性、太阳质量损失影响",
     "直接映射为三个候选方向"],
]

P8_EXAMPLE = (
    "q049 实例（V1 §1–2）：四条题干前提校正——理想牛顿力学下轨道不衰变；真实系统存在微弱耗散但时标极长；失稳主要源于混沌而非耗散；"
    "太阳演化（红巨星）是最相关的时间边界。配套变量表（a/e/i、Q、GW 功率、质量损失率）与「已有认识/争议/未知」清单，"
    "全部进入 D1–D3 的生成上下文。"
)

P9_TABLE = [
    ["每个缺口映射机制层解释：保守动力学失稳 / 微弱耗散累积 / 太阳演化边界", "避免缺口悬置、假设凭空而来"],
    ["强制三个方向＋显式差异维度（时间尺度/确定性/观测支持）", "防止同质重复输出"],
    ["每方向附支持依据、反对证据与不确定性", "防止单边论证"],
    ["差异维度＋可区分预测双约束；评审对 Direction 质量单独打分", "淘汰重复、空泛或不可检验输出"],
    ["核心陈述/依据/反证/可检验预测/不确定性五元组＋横向比较表", "结构一致才能逐项评审与筛选"],
]

P9_OUTPUT = (
    "每条候选方向包含：核心陈述、支持依据（带来源）、反对证据或替代解释、可区分预测、不确定性五要素，"
    "并附三方向横向比较（时间尺度、确定性、观测支持）；每条声明处理结果（选为主方向/降级/不选）及理由，处理结果随版本链保留。"
)

P10_TABLE = [
    ["前提校正后对象/范围与原题一致", "偏题方向在横向比较中剔除"],
    ["评审逐条回读来源并与原文比对", "Rasio 1996「Earth may well not survive」反向转述被发现并重写"],
    ["引用抽查表逐条判定 pass/fail", "V1 5 条仅 2 条通过→删错配源、修正 arXiv 号→V8 6/6"],
    ["判据必须可执行且物理正确", "「dE/dt＜10⁻²⁰ W」判据被判无效→改用 Peters 公式与 inspiral 时间比较"],
    ["三方向机制层面互斥检查", "D1 混沌 / D2 耗散 / D3 太阳演化保持机制可区分"],
    ["主方向＋补充＋稀有事件分层保留", "D3 选主、D2 降为补充、D1 限为稀有失稳（约 1% 概率归因已修正）"],
]

P10_RESULT = (
    "保留/降级/淘汰实例：D1 降级——概率归因与来源链不足；D2 不选为主——V1 把地球—太阳引力波功率误写为约 10⁻²⁰ W（实际约 200 W）且判据无效；"
    "D3 选主但重写 Rasio 转述。全部处理结果与被拒原因随 v1–v8 与各轮 review 保留，可回放。"
)

P11_TABLE = [
    ["耗散时标远超太阳寿命，轨道不会因该机制螺旋坠日；水星失稳为小概率稀有事件", "预测直接来自 D3 主线与 D1 限定"],
    ["N 体积分、广义相对论修正、太阳质量损失、潮汐模型、Monte Carlo——全部标 planned", "条件与假设检验手段一一对应"],
    ["Peters 圆轨道计算（executed：P=196.291 W、t=1.069e23 yr）→含 GR 的 N 体积分→潮汐与 MC 统计", "已执行项验证计算链路，planned 项待资源"],
    ["inspiral 时间＜太阳寿命→耗散主导；否则耗散仅为长期背景、混沌决定稀有失稳", "判据双向可判定，支持或反对均明确"],
    ["判据无效即回退重算（V1 判据被 Peters 替代）；计算未存 Artifact 时以输入/哈希/复算兜底", "防止无效判据误导后续实施"],
]

P11_BULLETS = [
    ("团队如何检查研究计划是否具体、可复核：",
     "独立评审验证判据物理正确性（曾发现 22 个数量级错误）；回执复算输出与 SHA-256 一致才认定 executed。"),
    ("研究计划中哪些内容已经具备执行条件，哪些仍是建议：",
     "executed 仅 Peters 地球—太阳计算；N 体/相对论/质量损失/潮汐/Monte Carlo 全部 planned，不冒充已执行。"),
    ("本作品如何避免用“进一步研究”等空泛表述代替研究步骤：",
     "每步写明输入、公式/命令、产物与停止条件；无法落地者标 planned，禁止以「进一步研究」替代。"),
]

P12_TABLE = [
    ["错配 DOI、错误 arXiv 号、反向转述",
     "review-v1 引用抽查 3 条 fail（有效率 2/5）",
     "删除错配来源、修正标识符、按原文限定重写",
     "V2 起引用 5/5；来源门此后全过"],
    ["定量错误与无效判据",
     "引力波功率错约 22 个数量级；原判据会把真实约 200 W 误判为显著",
     "Peters 公式＋完整输入＋独立复算替代",
     "V3 达 12/12；判据改为 inspiral 时间与太阳寿命比较"],
    ["传输失败与人工暂停",
     "上游 502/429 致无产物；q041–q045 人工暂停",
     "failed 留痕＋全新 Session 重试；stopped 单独记录",
     "全量 failed=0；恢复后每题唯一候选结论"],
]

P12_BULLETS = [
    ("哪些反馈由系统自动产生：",
     "verdict 与 findings、引用抽查 pass/fail、回执复算、终态机（completed/partial/waiting_human/failed）。"),
    ("哪些调整来自研究者或团队成员：",
     "q041–q045 暂停与恢复决定；伦理/招募/阈值类题保留 waiting_human 人工 Gate；scale-up 复审 GO/NO-GO。"),
    ("团队实际保留了哪些前后变化，便于说明反馈是否有效：",
     "v1–v8 全版本＋每轮 review＋失败 Session 记录＋关键文件 SHA-256；findings→修订→复评前后对照可回放。"),
]

P13_METHOD = (
    "Q001–Q125 每题由一个独立 Pi 作者 Session 运行同一轻量协议：读取 canonical 问题，输出证据边界、三个可区分 Direction、横向比较、研究计划、"
    "来源与候选结论；传输失败只重试一次，停止与错误 Session 原样保留。五个独立审计 Session 各复核连续 25 题，run.md frontmatter 独占 Project 终态。"
    "共同结果均为 Qwen 候选结论与 planned 研究方案，不是已验证发现；全量结果与代价见 P19。"
)

P13_CASE_BULLETS = [
    ("案例所属的科学领域及问题特点：",
     "q049 属天文学（太阳系动力学）；题干「轨道衰变终将坠日」含误导性前提，必须先校正再作答。"),
    ("团队选择它进行完整展示的实际理由：",
     "该题覆盖错误前提校正、来源核验、三机制比较、限定计算、独立评审、修订与直接回答对照全部环节。"),
    ("该案例能够展示本作品的哪些关键能力：",
     "错引发现与修复、22 个数量级定量纠错、planned/executed 分离、评审-修订闭环与回执复算。"),
    ("该案例表现不能代表哪些题目或条件：",
     "需付费全文、受限数据、伦理审批、实验设施或领域阈值裁决的题目（如 q021/q112/q098 中留待人工 Gate 的部分）。"),
]

P14_TABLE = [
    ["Why don't the orbits of planets decay and cause them to crash into each other?（题干断言：轨道会逐渐衰变，行星终将螺旋坠入太阳）"],
    ["太阳系八大行星轨道长期稳定性；关键变量：半长轴 a、偏心率 e、倾角 i、潮汐品质因子 Q、引力波功率、太阳质量损失率"],
    ["Gyr 尺度数值稳定（Laskar 1989/2004；Batygin & Laughlin 2008）；水星 Lyapunov 时间约 5 Myr；红巨星阶段将改变内行星轨道"],
    ["各机制时间尺度对比；混沌与耗散相对重要性；太阳质量损失对轨道的影响"],
    ["水星失稳概率 1–2% 或更高；外行星 ＞10 Gyr 超长期稳定性未知"],
]

P14_BULLETS = [
    ("第一轮使用了哪些证据和约束：",
     "同行评议/机构一手来源优先；精确主张须有可回读标识符；未执行的模拟与实验一律标 planned。"),
    ("第一轮要求系统生成什么内容：",
     "三个机制 Direction、横向比较、研究计划、来源记录与候选结论。"),
    ("第一轮采用了哪些主要模型或方法设置：",
     "contest-qwen/qwen3-max 独立作者 Session；4 次学术搜索、2 次 DOI 提取、25 次调用、非缓存输入 98,844 token。"),
    ("第一轮结果使用什么口径评价：",
     "独立评审六维 rubric（满分 12）＋全部来源逐条核验（V1 分母 5 条）。"),
]

P15_H = [
    ["H-01 · D1 N 体混沌可能导致低概率失稳",
     "Laskar 混沌扩散；水星 Lyapunov 约 5 Myr",
     "概率归因与来源链不完整",
     "长期积分给出的失稳率分布",
     "降级；概率归因与来源链需修正"],
    ["H-02 · D2 潮汐、引力波等微弱耗散长期累积",
     "潮汐理论与引力波辐射",
     "V1 误写地球—太阳引力波功率约 10⁻²⁰ W（实际约 200 W），判据 dE/dt＜10⁻²⁰ W 无效",
     "轨道衰变率的直接计算/测量",
     "不选；功率量级和计划判据错误"],
    ["H-03 · D3 太阳演化先于耗散决定内行星命运",
     "恒星演化理论；红巨星阶段轨道改变",
     "所引 Rasio 1996 结论被反向转述",
     "红巨星阶段内行星轨道改变的时点",
     "选为主方向；重写转述后保留"],
]

P15_PLAN = [
    ["含 GR 修正的 N 体长期积分（planned）", "区分保守稳定与混沌失稳，检验水星失稳概率"],
    ["潮汐与引力波耗散建模（planned）", "量化各机制时标，比较是否接近太阳寿命"],
    ["Monte Carlo 初始条件采样（planned）", "给出失稳概率区间，回答「低概率」是否成立"],
]

P16_TABLE = [
    ["关键来源不可用：Deienno/Nesvorný DOI 属另一论文",
     "review-v1 抽查：该 DOI 实际指向 Pires et al.（Icarus 246）",
     "D3 依据链断裂，关键论据失据",
     "删除错配来源；Lecar arXiv 号 0111602→0111600 修正并补齐记录"],
    ["关键结论反向转述：Rasio 原文「Earth may well not survive」被写成「可能幸存」",
     "原文逐句比对",
     "太阳演化论据方向整体反转",
     "按原文限定重写 D3 表述"],
    ["定量骨架失真：引力波功率错约 22 个数量级；判据无效",
     "Peters 公式独立复算约 200 W",
     "会把真实耗散误判为显著并误导计划实施",
     "以 Peters 公式＋完整输入＋独立复算替代；判据改为 inspiral 时间与太阳寿命比较"],
]

P16_BULLETS = [
    ("哪条证据、评价或人工意见触发了调整：",
     "review-v1 的五条独立 findings：错配 DOI、错误 arXiv 号、反向转述、22 个数量级错误、无效判据。"),
    ("第二轮增加、删除或改变了什么：",
     "删错源、修正标识符、重写反向转述、以 Peters 公式替换判据；V2 起引用 5/5。"),
    ("哪些内容保持不变，为什么：",
     "三方向机制框架与 planned/executed 边界不变——评审未否定框架本身，问题集中在来源与定量。"),
    ("团队预期第二轮在哪些方面改善：",
     "引用有效率与定量骨架；实际 V2 8/12（引用 5/5），继续修订至 V3 12/12。"),
]

P17_TABLE = [
    ["5 条来源仅 2/5 有效，含错配与错号", "6/6 全部可回读", "逐条核验＋删源修正", "来源门全过"],
    ["D3 含反向转述；D2 定量错误", "主线改为「保守近可积系统不自然衰变」为主、微弱耗散为补充、混沌为稀有失稳", "按原文限定＋复算重构", "主线方向正确且论证闭合"],
    ["无有效定量判据", "inspiral 时间与太阳寿命比较", "Peters 独立复算", "筛选判据可执行"],
    ["判据会把约 200 W 误判为显著", "门槛可执行：Peters 计算 executed（P=196.291 W、t=1.069e23 yr）", "完整输入＋独立复算", "reviewer 判 deliverable"],
    ["planned/executed 未显式分离", "唯一 executed＝Peters 计算，其余全部 planned", "边界显式标注", "结论限定成立"],
]

P17_BULLETS = [
    ("第二轮实际改善了什么：",
     "9/12→12/12；来源 2/5→6/6；判据物理有效；回执复算一致（review-v8 判 deliverable）。"),
    ("哪些方面没有改善或出现了新的代价：",
     "多轮检索/修订/审核代价——q049 链 920 次调用、非缓存输入约 6.28M token；限定计算脚本未存 Artifact（以输入/公式/命令/输出/哈希/复算留痕）。"),
    ("团队为什么在此时停止，或为什么仍需继续迭代：",
     "最终评审可交付＋回执一致＋剩余项全部 planned；约 1% 水星失稳概率的来源归因已修正，继续修订边际收益低。"),
]

P18_TABLE = [
    ["直接回答 attempt 2（实算近似对照）",
     "同题、同模型（qwen3-max）、同检索权限；各 5 次成功搜索＋1 次 write",
     "六维 rubric＋引用核验",
     "Workflow V1：9/12、来源 5 条（2/5）、3 方向、计划有但判据错误",
     "4/12；0 条显式来源；无 Direction、无计划",
     "calls 少 16.00%、非缓存输入多 14.65%，但长度仅 V1 的 48.05%——只近似计算量"],
    ["直接回答 attempt 6（长度近似对照）",
     "同上",
     "同上",
     "同上",
     "6/12；0 条显式来源；经七次 Crossref curl 检索",
     "长度 94.73%、calls 108.00%，但非缓存输入为 V1 的 11.968 倍——只近似长度"],
    ["对照结论",
     "—",
     "—",
     "最终版 12/12、来源 6/6",
     "两个对照均不可直接作为学术答案",
     "12/12 发生在独立评审、修订与限定计算之后，不能只归因于 Workflow"],
]

P18_BULLETS = [
    ("科学逻辑方面，本作品实际改善了什么：",
     "显式来源＋可执行判据＋planned/executed 边界，使结论可回读、可限定、可证伪。"),
    ("技术方法方面，本作品实际改善了什么：",
     "作者/评审分离、findings 驱动修订、回执复算，形成可审计闭环。"),
    ("结果表现方面，本作品实际改善了什么：",
     "同条件直接回答 4/12 与 6/12、0 来源；本作品 V1 9/12→最终 12/12、来源 6/6。"),
    ("没有改善的部分及增加的成本：",
     "全量未做同题重复运行，不声称跨采样稳健；五深度案例合计 3144 次调用、非缓存输入约 19.22M token；外部 benchmark 只证组件门，不作为科学结果。"),
]

P19_TABLE = [
    ["8（6.4%）", "轻量门槛全过（completed）", "候选结论 125/125；含 q038/q047/q048/q053 等"],
    ["117（93.6%）", "至少一项未过轻量门槛：来源层级/主张映射/结构/元数据/canonical 身份/planned 边界（partial）", "已有可提交、可回读结论，不等于未运行"],
    ["0", "waiting_human 与 failed 均为 0；5 个 stopped Session 是 q041–q045 人工暂停记录，不是 Project failed", "深度案例中 3 题因伦理/设施留待人工 Gate"],
    ["未开展", "全量未做同题重复运行", "不声称跨采样稳健性；深度案例内部多轮修订可复现"],
]

P19_FAIL_TABLE = [
    ["误导性题干前提", "部分题目题干含错误前提（如 q049「坠日」）", "生成时未强制前提校正会接受前提", "前提校正已写入协议并在全量轻量协议中执行"],
    ["错引与元数据拼接", "q021 正确 PMID 拼错误 DOI；q089 DOI 有效但题名/作者错配", "标识符可解析≠元数据正确", "来源门逐条回读＋回执复算可发现并修复"],
    ["评审过早通过／planned 写成 executed", "q112 10/12 时仍留错用标准与错误方程；review-v15 被独立回执否决", "reviewer 的 deliverable 不是免检信号", "分支级独立验收＋人工 Gate 兜底；继续修订至全过"],
]

P19_BULLET1 = "本作品未开展 125 题之外的泛化测试，故不设置该部分。"
P19_BULLET2 = "候选假设和研究计划仍需研究者审查，不等同于已经获得科学发现或完成真实验证。"

P20_TABLE = [
    ["GitHub LittleDrinks/research-world（main 分支持续维护至评审期）；cd research-world && docker compose up --build -d；服务 control:8095 / runtime:8098 / runner-controller:8096；测试：uv run pytest＋web npm test"],
    ["qwen3-max 经阿里云百炼 OpenAI 兼容端点；调用凭证截图【待用户补充：不泄露密钥的百炼调用凭证截图】；.env（apikey/baseurl）不入库"],
    ["evidence/contest-2026/all/q001–q125.md＋index.md＋5 份 audit-*.md；结果站走 GitHub Pages（push 即更新）；固定存档 tag：contest-2026-submission"],
    ["q049：project.json、v1–v8、review-v1–v8、receipt-v10、run.md（含对照 attempt 1–7 与显示投影）；q089/q021/q112/q098 同构证据链"],
    ["[DEPLOY-URL]（部署完成后回填真实可交互地址与示例请求/响应）；T2 Playwright 5 项验收（流式/刷新恢复/无内部泄露/纯文本轮/凭证失败可见）见 web/test-results/t2/；产品边界：研究工作流以 CLI 交付（Pi Session＋运行账本），Web 界面提供真实模型对话与图谱查看"],
    ["【待用户补充：≤10 分钟粗剪版，上传夸克网盘后回填链接】"],
]

# ---------------------------------------------------------------------------
# 填充
# ---------------------------------------------------------------------------

def unpagebreak_headings(doc):
    """模板 P2–P20 标题均强制分页；官方说明页写明“无需逐页机械对应”，
    为满足 ≤20 页限制，取消强制分页并给标题加 keep-with-next。"""
    from docx.oxml.ns import qn as _qn

    for p in doc.paragraphs:
        pPr = p._element.find(_qn("w:pPr"))
        if pPr is None:
            continue
        pbb = pPr.find(_qn("w:pageBreakBefore"))
        if pbb is not None:
            pPr.remove(pbb)
        if p.style is not None and p.style.name.startswith("Heading"):
            kwn = pPr.find(_qn("w:keepNext"))
            if kwn is None:
                pPr.append(pPr.makeelement(_qn("w:keepNext"), {}))


def main():
    doc = docx.Document(str(TEMPLATE))
    unpagebreak_headings(doc)

    # -- P1 ------------------------------------------------------------
    # 报名表截图占位框
    for p in doc.paragraphs:
        if p._element.findall(".//" + qn("w:drawing")):
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            fmt(p.add_run("【待用户补充】此处贴入盖章版报名表第一页、第二页截图（含个人信息，不入 git，提交时由用户手工插入）。"), 10.5, bold=True)
            break
    fill_table(doc, 0, 0, 1, [
        [PENDING + "：以挑战杯系统报名名称为准】"],
        ["Research World：面向 125 个科学问题的科学假设生成与研究计划设计智能体系统（建议稿）"],
        ["赛道一-方向1A-科学假设生成与研究计划设计（题目编号 XH-202619）"],
        [INTRO],
        [QWEN_NOTE],
        [PENDING + "：演示视频（≤10 分钟）上传夸克网盘后的分享链接，制作完成后回填】"],
    ], size=9.5)
    for prefix, text in CLAIMS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P2 ------------------------------------------------------------
    replace_para(doc, "[请用一段话说明", P2_PROBLEM)
    replace_para(doc, "[请用一段话概括", P2_WORK)
    insert_figure(doc, "[请插入本作品的总体思路图", "fig1-overview.png",
                  "图1 总体思路：从科学问题到候选结论的闭环（橙色为评审-修订回路）")
    for prefix, text in P2_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P3 ------------------------------------------------------------
    fill_table(doc, 1, 1, 1, P3_TABLE)
    replace_para(doc, "[请结合本作品实际方法说明", P3_LOGIC)

    # -- P4 ------------------------------------------------------------
    fill_table(doc, 2, 1, 0, P4_TABLE)
    for prefix, text in P4_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P5 ------------------------------------------------------------
    fill_table(doc, 4, 1, 0, P5_TABLE)
    for prefix, text in P5_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P6 ------------------------------------------------------------
    insert_figure(doc, "[请插入本作品实际架构图", "fig2-architecture.png",
                  "图2 系统总体架构：实线为已实现，虚线为规划中（ADR-0037）")
    fill_table(doc, 6, 1, 1, P6_TABLE)
    replace_para(doc, "[请说明本作品为什么不是一次性问答", P6_LOOP)

    # -- P7 ------------------------------------------------------------
    fill_table(doc, 7, 1, 1, P7_TABLE)
    insert_figure(doc, "[请插入一份真实的上下文结构示意", "fig3-context.png",
                  "图3 Qwen 上下文结构：作者侧分层拼接，评审侧隔离（红色为 findings 回流）")
    for prefix, text in P7_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P8 ------------------------------------------------------------
    fill_table(doc, 8, 1, 1, P8_TABLE)
    replace_para(doc, "[请选择一道实际测试题", P8_EXAMPLE)

    # -- P9 ------------------------------------------------------------
    fill_table(doc, 9, 1, 1, P9_TABLE)
    replace_para(doc, "[请说明本作品每条候选假设真实包含", P9_OUTPUT)

    # -- P10 -----------------------------------------------------------
    fill_table(doc, 10, 1, 1, P10_TABLE)
    replace_para(doc, "[请说明团队实际如何保留、合并、降级或淘汰", P10_RESULT)

    # -- P11 -----------------------------------------------------------
    fill_table(doc, 11, 1, 1, P11_TABLE)
    for prefix, text in P11_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P12 -----------------------------------------------------------
    insert_figure(doc, "[请插入本作品真实工作流程图", "fig4-workflow.png",
                  "图4 完整运行流程（q049 实链）：评审 findings 触发修订，回执复算收束")
    fill_table(doc, 12, 1, 0, P12_TABLE)
    for prefix, text in P12_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P13 -----------------------------------------------------------
    replace_para(doc, "[请说明本作品如何对官方125道科学问题进行全量测试", P13_METHOD)
    for prefix, text in P13_CASE_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P14 -----------------------------------------------------------
    fill_table(doc, 14, 1, 1, P14_TABLE)
    for prefix, text in P14_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P15 -----------------------------------------------------------
    fill_table(doc, 15, 1, 0, P15_H)
    fill_table(doc, 16, 1, 1, P15_PLAN)

    # -- P16 -----------------------------------------------------------
    fill_table(doc, 18, 1, 0, P16_TABLE)
    for prefix, text in P16_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P17 -----------------------------------------------------------
    fill_table(doc, 20, 1, 1, P17_TABLE)
    for prefix, text in P17_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P18 -----------------------------------------------------------
    fill_table(doc, 21, 1, 0, P18_TABLE)
    for prefix, text in P18_BULLETS:
        replace_para(doc, prefix, text, keep_bullet=prefix)

    # -- P19 -----------------------------------------------------------
    fill_table(doc, 23, 1, 1, P19_TABLE)
    insert_figure_after_table(doc, 23, "fig5-results.png",
                              "图5 125 题结果可视化：(a) 终态分布；(b) 深度案例 rubric；(c) 全量成本（对数轴）")
    fill_table(doc, 24, 1, 0, P19_FAIL_TABLE)
    replace_para(doc, "如团队实际开展了125题之外的泛化测试", P19_BULLET1)
    replace_para(doc, "请明确候选假设和研究计划仍需研究者审查", P19_BULLET2)

    # -- P20 -----------------------------------------------------------
    fill_table(doc, 25, 1, 1, P20_TABLE)

    clear_underscore_lines(doc)
    # 删除文末空段，避免尾部空白页
    body = doc.element.body
    for child in list(body.iterchildren())[::-1]:
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p") and not "".join(child.itertext()).strip():
            body.remove(child)
        else:
            break
    doc.save(str(OUT))
    print(f"saved {OUT}")


if __name__ == "__main__":
    main()
